#!/usr/bin/env python3
"""
Convert thesis.md to a properly formatted academic DOCX.
Uses pandoc for base conversion + python-docx post-processing.
Embeds all figures (Ch3 + Ch4), preserves equations, formats tables.
"""

import re
import subprocess
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

THESIS_DIR = Path(__file__).parent
ASSETS_DIR = THESIS_DIR / "assets"
OUTPUT_DOCX = THESIS_DIR / "thesis.docx"
REFERENCE_DOCX = THESIS_DIR / "reference_docx.docx"
PREPROCESSED_MD = THESIS_DIR / "_thesis_pandoc_ready.md"


def create_reference_docx():
    """Create a reference.docx with academic thesis styles for pandoc."""
    doc = Document()

    # Page setup: A4, academic margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # Title
    s = doc.styles["Title"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(18)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(24)

    # Subtitle
    s = doc.styles["Subtitle"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(14)
    s.font.bold = False
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(12)

    # Headings
    for level, (sz, bld, sb, sa) in {
        1: (16, True, 24, 12),
        2: (14, True, 18, 8),
        3: (12, True, 12, 6),
        4: (12, True, 12, 6),
    }.items():
        s = doc.styles[f"Heading {level}"]
        s.font.name = "Times New Roman"
        s.font.size = Pt(sz)
        s.font.bold = bld
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(sb)
        s.paragraph_format.space_after = Pt(sa)
        s.paragraph_format.line_spacing = 1.5

    # Caption
    s = doc.styles["Caption"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(10)
    s.font.italic = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(6)
    s.paragraph_format.space_after = Pt(12)

    doc.save(str(REFERENCE_DOCX))
    print("  [OK] reference_docx.docx created")


def preprocess_thesis_md():
    """
    Preprocess thesis.md for pandoc conversion:
    - Add YAML metadata
    - Embed figure images inline with captions
    """
    raw = (THESIS_DIR / "thesis.md").read_text(encoding="utf-8")

    # YAML header
    yaml = """---
title: "Detection of Social Media Deepfake Contents Using Deep Learning Algorithm"
author: "Olamijulo Israel D"
date: "February 2026"
documentclass: report
classoption:
  - 12pt
  - a4paper
  - oneside
geometry:
  - top=2.54cm
  - bottom=2.54cm
  - left=3.17cm
  - right=2.54cm
fontsize: 12pt
linestretch: 1.5
numbersections: false
toc: true
toc-depth: 3
header-includes:
  - \\usepackage{{booktabs}}
  - \\usepackage{{longtable}}
  - \\usepackage{{graphicx}}
  - \\usepackage{{float}}
  - \\usepackage{{caption}}
  - \\captionsetup{{font=small, labelfont=bf}}
---

"""
    content = yaml + raw

    # ── Embed figures inline ─────────────────────────────────
    # Each tuple: (search_text, replacement_text)
    # Order matters — longer/more-specific strings first for Ch4 figures
    # since some Ch3 patterns could be substrings.

    figure_embeddings = [
        # ================================================================
        # CHAPTER 3: METHODOLOGY FIGURES (Figures 3.1–3.9)
        # ================================================================

        # Figure 3.1 — System Architecture (after §3.1 phases list)
        (
            "6. **Deployment:** Exposing trained models via a FastAPI inference service.\n\n## 3.2 Datasets",
            (
                "6. **Deployment:** Exposing trained models via a FastAPI inference service.\n\n"
                "![Figure 3.1: Overview of the deepfake detection system architecture showing the end-to-end pipeline "
                "from data ingestion through preprocessing, model inference, and result delivery. "
                "](assets/ch3_system_architecture.png){#fig:system_architecture width=100%}\n\n"
                "## 3.2 Datasets"
            ),
        ),

        # Figure 3.2 — Preprocessing Pipeline (after §3.3 steps list)
        (
            "4. **Normalization:** Resize to model input size (299x299 for XceptionNet, 224x224 for EfficientNet), normalize to [0, 1] or ImageNet statistics.\n\n## 3.4 Model Architecture",
            (
                "4. **Normalization:** Resize to model input size (299x299 for XceptionNet, 224x224 for EfficientNet), normalize to [0, 1] or ImageNet statistics.\n\n"
                "![Figure 3.2: Preprocessing pipeline illustrating the four-stage process: frame extraction, "
                "face detection, face cropping, and normalisation. "
                "](assets/ch3_preprocessing_pipeline.png){#fig:preprocessing_pipeline width=100%}\n\n"
                "## 3.4 Model Architecture"
            ),
        ),

        # Figure 3.3 — XceptionNet Architecture (after §3.4 XceptionNet specs)
        (
            "- **Input Resolution:** 299×299\n\n### EfficientNet-B0 (Secondary)",
            (
                "- **Input Resolution:** 299×299\n\n"
                "![Figure 3.3: XceptionNet architecture with depthwise separable convolutions used as the primary "
                "detection model. Transfer learning from ImageNet with fine-tuning on FaceForensics++. "
                "](assets/ch3_xceptionnet_architecture.png){#fig:xceptionnet_architecture width=100%}\n\n"
                "### EfficientNet-B0 (Secondary)"
            ),
        ),

        # Figure 3.4 — EfficientNet-B0 Architecture (after §3.4 EfficientNet specs)
        (
            "- **Input Resolution:** 224×224\n\n## 3.5 Training Configuration",
            (
                "- **Input Resolution:** 224×224\n\n"
                "![Figure 3.4: EfficientNet-B0 architecture with compound scaling used as the secondary "
                "detection model. Lightweight design with 4.01M parameters. "
                "](assets/ch3_efficientnet_architecture.png){#fig:efficientnet_architecture width=100%}\n\n"
                "## 3.5 Training Configuration"
            ),
        ),

        # Figure 3.5 — Transfer Learning (after §3.4 intro, before XceptionNet subsection)
        (
            "## 3.4 Model Architecture\n\n### XceptionNet (Primary)",
            (
                "## 3.4 Model Architecture\n\n"
                "![Figure 3.5: Transfer learning process: pre-trained ImageNet weights are used to initialise "
                "the CNN backbone, then the classification head is fine-tuned on the deepfake dataset. "
                "](assets/ch3_transfer_learning.png){#fig:transfer_learning width=100%}\n\n"
                "### XceptionNet (Primary)"
            ),
        ),

        # Figure 3.6 — Training Pipeline (after §3.5 hyperparameter table)
        (
            "| Persistent Workers | Enabled |\n\n### 3.5.1 Loss Function",
            (
                "| Persistent Workers | Enabled |\n\n"
                "![Figure 3.6: Training pipeline showing the end-to-end process from dataset loading through "
                "augmentation, model training with label smoothing, and checkpoint management. "
                "](assets/ch3_training_pipeline.png){#fig:training_pipeline width=100%}\n\n"
                "### 3.5.1 Loss Function"
            ),
        ),

        # Figure 3.7 — Evaluation Pipeline (after §3.6 ROC-AUC explanation)
        (
            "An AUC of 1.0 represents a perfect classifier, while an AUC of 0.5 represents random guessing.\n\n## 3.7 Video-Level Aggregation",
            (
                "An AUC of 1.0 represents a perfect classifier, while an AUC of 0.5 represents random guessing.\n\n"
                "![Figure 3.7: Evaluation pipeline showing the metrics computation process including "
                "confusion matrix derivation, per-class metrics, and ROC-AUC calculation. "
                "](assets/ch3_evaluation_pipeline.png){#fig:evaluation_pipeline width=100%}\n\n"
                "## 3.7 Video-Level Aggregation"
            ),
        ),

        # Figure 3.8 — Video Aggregation (after §3.7 methods list)
        (
            "3. **Confidence Weighting:** Weighted average by prediction confidence\n\n## 3.8 Tools and Technologies",
            (
                "3. **Confidence Weighting:** Weighted average by prediction confidence\n\n"
                "![Figure 3.8: Video-level aggregation strategies: mean probability averaging, majority voting, "
                "and confidence-weighted combining of frame-level predictions. "
                "](assets/ch3_video_aggregation.png){#fig:video_aggregation width=100%}\n\n"
                "## 3.8 Tools and Technologies"
            ),
        ),

        # Figure 3.9 — Inference Pipeline (after §3.8 tools table)
        (
            "| GradCAM | Model explainability |\n\n## 3.9 Ethical Considerations",
            (
                "| GradCAM | Model explainability |\n\n"
                "![Figure 3.9: Inference pipeline for real-time deepfake detection: video input, frame extraction, "
                "face detection, model prediction, and aggregation to final verdict. "
                "](assets/ch3_inference_pipeline.png){#fig:inference_pipeline width=100%}\n\n"
                "## 3.9 Ethical Considerations"
            ),
        ),

        # ================================================================
        # CHAPTER 4: RESULTS & DISCUSSION FIGURES (Figures 4.1–4.6)
        # ================================================================

        # Figure 4.1 — XceptionNet training curves (after §4.2.1 training table)
        (
            "| 30 | 0.0143 | 99.68% | 0.0164 | 99.61% | 0.000125 | 99.3 |\n\n### 4.2.2",
            (
                "| 30 | 0.0143 | 99.68% | 0.0164 | 99.61% | 0.000125 | 99.3 |\n\n"
                "![XceptionNet training loss and accuracy curves over 30 epochs. "
                "*Left:* Training and validation loss. *Right:* Training and validation accuracy. "
                "](assets/XceptionNet_training_curves.png){#fig:xception_training width=100%}\n\n"
                "### 4.2.2"
            ),
        ),

        # Figure 4.2 — EfficientNet training curves (after §4.2.2 training table)
        (
            "| 29 | 0.0092 | 99.78% | 0.0082 | 99.85% | 0.000125 | 170.7 |\n\n## 4.3",
            (
                "| 29 | 0.0092 | 99.78% | 0.0082 | 99.85% | 0.000125 | 170.7 |\n\n"
                "![EfficientNet-B0 training loss and accuracy curves over 29 epochs. "
                "*Left:* Training and validation loss. *Right:* Training and validation accuracy. "
                "](assets/EfficientNet-B0_training_curves.png){#fig:efficientnet_training width=100%}\n\n"
                "## 4.3"
            ),
        ),

        # Figure 4.3 — XceptionNet confusion matrix (after §4.4.3 analysis bullets)
        (
            "- **False Positive Rate:** 0.57%\n- **False Negative Rate:** 1.00%\n\n**Table 4.10",
            (
                "- **False Positive Rate:** 0.57%\n- **False Negative Rate:** 1.00%\n\n"
                "![XceptionNet confusion matrix on the test set (5,965 samples). "
                "Diagonal elements represent correct predictions; off-diagonal elements represent misclassifications. "
                "](assets/XceptionNet_confusion_matrix.png){#fig:xception_cm width=70%}\n\n"
                "**Table 4.10"
            ),
        ),

        # Figure 4.4 — EfficientNet confusion matrix (after §4.4.3 analysis bullets)
        (
            "- **False Positive Rate:** 0.27%\n- **False Negative Rate:** 0.17%\n\n**Analysis:**",
            (
                "- **False Positive Rate:** 0.27%\n- **False Negative Rate:** 0.17%\n\n"
                "![EfficientNet-B0 confusion matrix on the test set (5,965 samples). "
                "Diagonal elements represent correct predictions; off-diagonal elements represent misclassifications. "
                "](assets/EfficientNet-B0_confusion_matrix.png){#fig:efficientnet_cm width=70%}\n\n"
                "**Analysis:**"
            ),
        ),

        # Figure 4.5 — ROC comparison + Figure 4.6 — Model comparison (after §4.5 trade-off)
        (
            "- Both models are suitable for real-time deployment on modern hardware\n\n## 4.6 Explainability Analysis",
            (
                "- Both models are suitable for real-time deployment on modern hardware\n\n"
                "![ROC curves comparison between XceptionNet (blue) and EfficientNet-B0 (orange). "
                "Both models achieve near-perfect AUC scores, with EfficientNet-B0 at 0.9997 and XceptionNet at 0.9992. "
                "](assets/roc_comparison.png){#fig:roc_comparison width=85%}\n\n"
                "![Comparative bar charts of key performance metrics. "
                "EfficientNet-B0 outperforms XceptionNet in accuracy, F1-score, and ROC-AUC while using 4.25x fewer parameters. "
                "](assets/model_comparison.png){#fig:model_comparison width=100%}\n\n"
                "## 4.6 Explainability Analysis"
            ),
        ),
    ]

    inserted = 0
    total = len(figure_embeddings)
    for search, replacement in figure_embeddings:
        if search in content:
            content = content.replace(search, replacement, 1)
            inserted += 1
            print(f"  [+] Embedded figure {inserted}/{total}")
        else:
            print(f"  [!] Could not find insertion point for figure ({search[:60]}...)")

    # Write preprocessed file
    PREPROCESSED_MD.write_text(content, encoding="utf-8")
    print(f"  [OK] Preprocessed markdown ({len(content)} chars, {inserted} figures embedded)")


def run_pandoc():
    """Run pandoc to convert preprocessed markdown to docx."""
    cmd = [
        "pandoc",
        "_thesis_pandoc_ready.md",
        "-o", "thesis.docx",
        "--toc",
        "--toc-depth=3",
        "--number-sections",
        "--wrap=none",
        "-f", "markdown+tex_math_dollars+pipe_tables+grid_tables+multiline_tables+implicit_figures",
        "-t", "docx",
    ]

    print(f"  [..] Running pandoc...")
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=str(THESIS_DIR))

    if result.returncode != 0:
        print(f"  [ERR] pandoc failed:\n{result.stderr}")
        return False

    if result.stderr:
        for line in result.stderr.strip().split("\n")[:5]:
            print(f"  [WARN] {line}")

    size_kb = OUTPUT_DOCX.stat().st_size / 1024
    print(f"  [OK] pandoc produced {OUTPUT_DOCX.name} ({size_kb:.0f} KB)")
    return True


def post_process_docx():
    """Post-process: verify elements, fix formatting."""
    if not OUTPUT_DOCX.exists():
        print("  [ERR] DOCX not found")
        return False

    doc = Document(str(OUTPUT_DOCX))

    # Enforce page layout
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)

    # Enforce Times New Roman 12pt on all body paragraphs
    for para in doc.paragraphs:
        style_name = para.style.name
        if style_name == "Normal" or style_name.startswith("List"):
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

    # Count embedded images
    images = [r for r in doc.part.rels.values() if "image" in r.reltype]
    tables = len(doc.tables)

    print(f"  [OK] Document statistics:")
    print(f"       Tables: {tables}")
    print(f"       Embedded images: {len(images)}")
    print(f"       Paragraphs: {len(doc.paragraphs)}")

    for img in images:
        print(f"         - {img.target_ref}")

    doc.save(str(OUTPUT_DOCX))
    return True


def main():
    print("=" * 60)
    print("  THESIS MARKDOWN -> DOCX CONVERSION")
    print("=" * 60)

    print("\n[1/4] Creating reference.docx...")
    create_reference_docx()

    print("\n[2/4] Preprocessing thesis.md...")
    preprocess_thesis_md()

    print("\n[3/4] Running pandoc...")
    if not run_pandoc():
        print("\n[FATAL] pandoc conversion failed.")
        return

    print("\n[4/4] Post-processing DOCX...")
    post_process_docx()

    # Cleanup temporary files
    if PREPROCESSED_MD.exists():
        PREPROCESSED_MD.unlink()
    if REFERENCE_DOCX.exists():
        REFERENCE_DOCX.unlink()

    print("\n" + "=" * 60)
    print(f"  DONE -> {OUTPUT_DOCX}")
    print("=" * 60)


if __name__ == "__main__":
    main()
